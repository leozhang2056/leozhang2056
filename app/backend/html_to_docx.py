from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from bs4 import BeautifulSoup, NavigableString, Tag, Comment
import yaml
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --- Constants for Unification ---
BASE_FONT_SIZE = 10.5
BASE_LINE_SPACING = 1.15
BASE_SPACE_AFTER = 4.0
# Spacing constants (use these to tune global rhythm)
SUMMARY_SPACE_AFTER = BASE_SPACE_AFTER * 2.0
SECTION_SPACE_BEFORE = 8.0
SECTION_SPACE_AFTER = BASE_SPACE_AFTER
CONTACT_SPACE_AFTER = BASE_SPACE_AFTER * 2.0
LIST_SPACE_AFTER = 1.5
SKILL_SPACE_AFTER = 1.5
EDU_TAB_SPACE_AFTER = 2.0
EDU_DETAIL_SPACE_AFTER = 3.0
JOB_HEADER_SPACE_AFTER = 2.0
JOB_ROLE_SPACE_AFTER = 1.5
TECH_SPACE_AFTER = 2.0

PDF_TITLE_BLUE = RGBColor(0x1A, 0x4D, 0x96)
PDF_LINK_BLUE = RGBColor(0x1A, 0x4A, 0x8A)
PDF_TEXT = RGBColor(0x11, 0x11, 0x11)
PDF_MUTED = RGBColor(0x44, 0x44, 0x44)
PDF_DATE = RGBColor(0x77, 0x77, 0x77)


# runtime-overridable values (can be overridden by app/backend/docx_style.yaml)
FONT_NAME = 'Roboto'
TAB_POS = 6.6


def _hex_to_rgb(hexstr: str) -> Optional[RGBColor]:
    if not hexstr:
        return None
    s = hexstr.strip().lstrip('#')
    if len(s) != 6:
        return None
    try:
        r = int(s[0:2], 16)
        g = int(s[2:4], 16)
        b = int(s[4:6], 16)
        return RGBColor(r, g, b)
    except Exception:
        return None


# Try to load overrides from YAML file in same folder
try:
    STYLE_PATH = Path(__file__).parent / 'docx_style.yaml'
    if STYLE_PATH.exists():
        cfg = yaml.safe_load(STYLE_PATH.read_text(encoding='utf-8')) or {}
        BASE_FONT_SIZE = float(cfg.get('base_font_size', BASE_FONT_SIZE))
        BASE_LINE_SPACING = float(cfg.get('base_line_spacing', BASE_LINE_SPACING))
        BASE_SPACE_AFTER = float(cfg.get('base_space_after', BASE_SPACE_AFTER))

        SUMMARY_SPACE_AFTER = float(cfg.get('summary_space_after', SUMMARY_SPACE_AFTER))
        SECTION_SPACE_BEFORE = float(cfg.get('section_space_before', SECTION_SPACE_BEFORE))
        SECTION_SPACE_AFTER = float(cfg.get('section_space_after', SECTION_SPACE_AFTER))
        CONTACT_SPACE_AFTER = float(cfg.get('contact_space_after', CONTACT_SPACE_AFTER))
        LIST_SPACE_AFTER = float(cfg.get('list_space_after', LIST_SPACE_AFTER))
        SKILL_SPACE_AFTER = float(cfg.get('skill_space_after', SKILL_SPACE_AFTER))
        EDU_TAB_SPACE_AFTER = float(cfg.get('edu_tab_space_after', EDU_TAB_SPACE_AFTER))
        EDU_DETAIL_SPACE_AFTER = float(cfg.get('edu_detail_space_after', EDU_DETAIL_SPACE_AFTER))
        JOB_HEADER_SPACE_AFTER = float(cfg.get('job_header_space_after', JOB_HEADER_SPACE_AFTER))
        JOB_ROLE_SPACE_AFTER = float(cfg.get('job_role_space_after', JOB_ROLE_SPACE_AFTER))
        TECH_SPACE_AFTER = float(cfg.get('tech_space_after', TECH_SPACE_AFTER))

        FONT_NAME = cfg.get('font_name', FONT_NAME)
        TAB_POS = float(cfg.get('tab_pos', TAB_POS))
        # element sizes
        HEADER_NAME_SIZE = float(cfg.get('header_name_size', 21))
        CONTACT_SIZE = float(cfg.get('contact_size', 9))
        SECTION_TITLE_SIZE = float(cfg.get('section_title_size', 12.7))
        SECTION_TITLE_LETTER_SPACING = float(cfg.get('section_title_letter_spacing', 0.7))

        # colors
        v = cfg.get('pdf_title_blue')
        if v:
            c = _hex_to_rgb(v)
            if c:
                PDF_TITLE_BLUE = c
        v = cfg.get('pdf_link_blue')
        if v:
            c = _hex_to_rgb(v)
            if c:
                PDF_LINK_BLUE = c
        v = cfg.get('pdf_text')
        if v:
            c = _hex_to_rgb(v)
            if c:
                PDF_TEXT = c
        v = cfg.get('pdf_muted')
        if v:
            c = _hex_to_rgb(v)
            if c:
                PDF_MUTED = c
        v = cfg.get('pdf_date')
        if v:
            c = _hex_to_rgb(v)
            if c:
                PDF_DATE = c
except Exception:
    pass

# fallback element sizes
HEADER_NAME_SIZE = globals().get('HEADER_NAME_SIZE', 21)
CONTACT_SIZE = globals().get('CONTACT_SIZE', 9)
SECTION_TITLE_SIZE = globals().get('SECTION_TITLE_SIZE', 12.7)
SECTION_TITLE_LETTER_SPACING = globals().get('SECTION_TITLE_LETTER_SPACING', 0.7)


def _set_document_defaults(document):
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    normal = document.styles['Normal']
    try:
        normal.font.name = FONT_NAME
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    except Exception:
        normal.font.name = 'Arial'
    normal.font.size = Pt(BASE_FONT_SIZE)

    fmt = normal.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(BASE_SPACE_AFTER)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _format_paragraph(paragraph, *, size: float = BASE_FONT_SIZE, bold: bool = False, italic: bool = False,
                      color: Optional[RGBColor] = None,
                      align: Optional[WD_PARAGRAPH_ALIGNMENT] = None,
                      space_before: float = 0,
                      space_after: float = BASE_SPACE_AFTER):
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for run in paragraph.runs:
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        try:
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        except Exception:
            pass


def _style_paragraph(paragraph, **kwargs):
    # Backward compatibility helper
    _format_paragraph(paragraph, **kwargs)


def _set_paragraph_bottom_border(paragraph, *, color: str = 'D9D9D9', size: str = '6') -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn('w:pBdr'))
    if p_bdr is None:
        p_bdr = OxmlElement('w:pBdr')
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn('w:bottom'))
    if bottom is None:
        bottom = OxmlElement('w:bottom')
        p_bdr.append(bottom)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)


def _set_run_small_caps(run) -> None:
    """Apply Word small-caps formatting to a run (matches HTML font-variant: small-caps)."""
    r_pr = run._element.get_or_add_rPr()
    existing = r_pr.find(qn('w:smallCaps'))
    if existing is None:
        small_caps = OxmlElement('w:smallCaps')
        small_caps.set(qn('w:val'), '1')
        r_pr.append(small_caps)


def _add_hyperlink(paragraph, url: str, text: str, *,
                   size: float = BASE_FONT_SIZE,
                   color: Optional[RGBColor] = None,
                   bold: bool = False) -> None:
    """Insert a real clickable hyperlink into a paragraph (matches PDF <a href>)."""
    if not url or not text:
        return
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), FONT_NAME)
    rfonts.set(qn('w:hAnsi'), FONT_NAME)
    rfonts.set(qn('w:eastAsia'), FONT_NAME)
    r_pr.append(rfonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    r_pr.append(sz)
    if bold:
        b = OxmlElement('w:b')
        r_pr.append(b)
    if color is not None:
        clr = OxmlElement('w:color')
        clr.set(qn('w:val'), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
        r_pr.append(clr)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'none')
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_table_borders_none(table) -> None:
    """Remove all visible borders from a table (for layout-only tables)."""
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        borders.append(b)
    tbl_pr.append(borders)


def _add_borderless_table_row(document, left_runs, right_runs, *,
                              left_size: float = BASE_FONT_SIZE,
                              right_size: float = BASE_FONT_SIZE - 0.5,
                              right_width_pct: float = 18.0,
                              space_after: float = 0.0,
                              right_align: str = 'right'):
    """
    Build a 1-row, 2-col borderless table where the right cell auto-sizes
    to a percentage of the page width. Used for lc-row (title + date).
    left_runs / right_runs are lists of (text, kwargs) tuples that get appended
    to the respective cell paragraph. kwargs can include 'bold', 'italic',
    'color', 'underline', 'hyperlink' (url).
    """
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    _set_table_borders_none(table)
    page_width = document.sections[0].page_width or Inches(8.5)
    left_margin = document.sections[0].left_margin or Inches(0.45)
    right_margin = document.sections[0].right_margin or Inches(0.45)
    usable = page_width - left_margin - right_margin
    right_w = int(usable * (right_width_pct / 100.0))
    left_w = int(usable - right_w)
    table.columns[0].width = left_w
    table.columns[1].width = right_w
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    left_cell, right_cell = table.rows[0].cells
    for text, kwargs in left_runs:
        if kwargs.get('hyperlink'):
            _add_hyperlink(left_cell.paragraphs[0], kwargs['hyperlink'], text,
                           size=kwargs.get('size', left_size),
                           color=kwargs.get('color'),
                           bold=kwargs.get('bold', False))
        else:
            r = left_cell.paragraphs[0].add_run(text)
            r.font.size = Pt(kwargs.get('size', left_size))
            r.font.name = FONT_NAME
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            if kwargs.get('bold'):
                r.bold = True
            if kwargs.get('italic'):
                r.italic = True
            if kwargs.get('color') is not None:
                r.font.color.rgb = kwargs['color']
    for text, kwargs in right_runs:
        if kwargs.get('hyperlink'):
            _add_hyperlink(right_cell.paragraphs[0], kwargs['hyperlink'], text,
                           size=kwargs.get('size', right_size),
                           color=kwargs.get('color'),
                           bold=kwargs.get('bold', False))
        else:
            r = right_cell.paragraphs[0].add_run(text)
            r.font.size = Pt(kwargs.get('size', right_size))
            r.font.name = FONT_NAME
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            if kwargs.get('bold'):
                r.bold = True
            if kwargs.get('italic'):
                r.italic = True
            if kwargs.get('color') is not None:
                r.font.color.rgb = kwargs['color']
    right_cell.paragraphs[0].alignment = {
        'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
        'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
        'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
    }.get(right_align, WD_PARAGRAPH_ALIGNMENT.RIGHT)
    return table


def _add_tabbed_line(document, left: str, right: str, *,
                     left_size: float = BASE_FONT_SIZE,
                     right_size: float = BASE_FONT_SIZE - 0.5,
                     left_bold: bool = False, right_bold: bool = False,
                     left_italic: bool = False, right_italic: bool = False,
                     left_color: Optional[RGBColor] = None, right_color: Optional[RGBColor] = None,
                     space_after: float = 2.0, tab_pos: float = None,
                     keep_with_next: bool = False):
    paragraph = document.add_paragraph()
    if tab_pos is None:
        tp = TAB_POS
    else:
        tp = tab_pos
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(tp), WD_TAB_ALIGNMENT.RIGHT)
    _format_paragraph(paragraph, size=left_size, space_after=space_after)
    if keep_with_next:
        paragraph.paragraph_format.keep_with_next = True

    left_run = paragraph.add_run(left)
    left_run.font.size = Pt(left_size)
    left_run.bold = left_bold
    left_run.italic = left_italic
    if left_color is not None:
        left_run.font.color.rgb = left_color
    try:
        left_run.font.name = FONT_NAME
        left_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    except Exception:
        pass

    paragraph.add_run('\t')

    right_run = paragraph.add_run(right)
    right_run.font.size = Pt(right_size)
    right_run.bold = right_bold
    right_run.italic = right_italic
    if right_color is not None:
        right_run.font.color.rgb = right_color
    try:
        right_run.font.name = FONT_NAME
        right_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    except Exception:
        pass
    return paragraph


def _add_text_paragraph(document, text: str, **kwargs):
    text = (text or '').strip()
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    _format_paragraph(paragraph, **kwargs)


def _append_inline(node, paragraph, size: float = BASE_FONT_SIZE, color: Optional[RGBColor] = PDF_TEXT):
    if isinstance(node, Comment):
        return
    if isinstance(node, NavigableString):
        # Collapse newlines and runs of whitespace into a single space — the source
        # HTML is multi-line and would otherwise leak \n into the rendered text.
        text = re.sub(r"\s+", " ", str(node))
        if text and text != " ":
            run = paragraph.add_run(text)
            run.font.size = Pt(size)
            run.font.color.rgb = color
            try:
                run.font.name = FONT_NAME
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            except Exception:
                pass
        elif text == " ":
            run = paragraph.add_run(" ")
            run.font.size = Pt(size)
            run.font.color.rgb = color
        return

    if not isinstance(node, Tag):
        return

    name = node.name.lower()
    if name == 'br':
        paragraph.add_run().add_break()
        return

    if name in {'strong', 'b'}:
        for child in node.children:
            if isinstance(child, NavigableString):
                if str(child):
                    run = paragraph.add_run(str(child))
                    run.bold = True
                    run.font.size = Pt(size)
                    run.font.color.rgb = color
                    try:
                        run.font.name = FONT_NAME
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                    except Exception:
                        pass
            else:
                _append_inline(child, paragraph, size=size, color=color)
        return

    if name in {'em', 'i'}:
        for child in node.children:
            if isinstance(child, NavigableString):
                if str(child):
                    run = paragraph.add_run(str(child))
                    run.italic = True
                    run.font.size = Pt(size)
                    try:
                        run.font.name = FONT_NAME
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                    except Exception:
                        pass
                    run.font.color.rgb = color
            else:
                _append_inline(child, paragraph, size=size, color=color)
        return

    if name == 'a':
        text = re.sub(r"\s+", " ", node.get_text(' ', strip=False))
        href = (node.get('href') or '').strip()
        if text and href:
            _add_hyperlink(paragraph, href, text, size=size, color=PDF_LINK_BLUE)
        elif text:
            run = paragraph.add_run(text)
            run.font.size = Pt(size)
            run.underline = True
            run.font.color.rgb = PDF_LINK_BLUE
            try:
                run.font.name = FONT_NAME
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            except Exception:
                pass
        return

    for child in node.children:
        _append_inline(child, paragraph, size=size, color=color)


def _render_inline_container(document, node: Tag, **kwargs):
    paragraph = document.add_paragraph()
    size = kwargs.get('size', BASE_FONT_SIZE)
    color = kwargs.get('color', PDF_TEXT)
    for child in node.children:
        _append_inline(child, paragraph, size=size, color=color)
    # Trim stray leading/trailing space from the inline content (HTML inter-tag
    # whitespace is harmless in browser rendering but shows up in DOCX text).
    if paragraph.runs:
        first = paragraph.runs[0]
        if first.text.startswith(' '):
            first.text = first.text.lstrip(' ')
            if not first.text:
                paragraph._p.remove(first._r)
        if paragraph.runs:
            last = paragraph.runs[-1]
            if last.text.endswith(' '):
                last.text = last.text.rstrip(' ')
                if not last.text:
                    paragraph._p.remove(last._r)
    _format_paragraph(paragraph, **kwargs)


def _render_section_title(document, text: str):
    paragraph = document.add_paragraph()
    # Use original case + small-caps (matches PDF font-variant: small-caps).
    # The first letter still renders larger per .section-title::first-letter CSS,
    # which is close enough visually to the all-caps UPPER variant.
    run = paragraph.add_run(text.strip())
    run.bold = True
    run.font.size = Pt(SECTION_TITLE_SIZE)
    _set_run_small_caps(run)
    try:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    except Exception:
        pass
    _format_paragraph(
        paragraph,
        size=SECTION_TITLE_SIZE,
        bold=True,
        color=PDF_TITLE_BLUE,
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        space_before=SECTION_SPACE_BEFORE,
        space_after=SECTION_SPACE_AFTER,
    )
    _set_paragraph_bottom_border(paragraph)
    # Keep the section title glued to the first paragraph of its content so the
    # title doesn't end up alone at the bottom of a page.
    paragraph.paragraph_format.keep_with_next = True


def _render_header(document, node: Tag):
    name_node = node.select_one('.cv-name')
    contact_node = node.select_one('.cv-contact')

    if name_node:
        _add_text_paragraph(
            document,
            name_node.get_text(' ', strip=True),
            size=HEADER_NAME_SIZE,
            bold=True,
            color=PDF_TEXT,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            space_after=CONTACT_SPACE_AFTER,
        )
    if contact_node:
        _render_inline_container(
            document,
            contact_node,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            size=CONTACT_SIZE,
            color=PDF_LINK_BLUE,
            space_after=CONTACT_SPACE_AFTER,
        )


def _render_list(document, node: Tag):
    items = node.find_all('li', recursive=False)
    last_idx = len(items) - 1
    for i, li in enumerate(items):
        paragraph = document.add_paragraph(style='List Bullet')
        _append_inline(li, paragraph, size=BASE_FONT_SIZE)
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.first_line_indent = Inches(-0.14)
        _format_paragraph(paragraph, size=BASE_FONT_SIZE, space_after=LIST_SPACE_AFTER, color=PDF_TEXT)
        # Keep each bullet on one page (no mid-bullet page break), and chain
        # consecutive bullets so the whole list prefers to stay together.
        paragraph.paragraph_format.keep_together = True
        if i < last_idx:
            paragraph.paragraph_format.keep_with_next = True


def _render_skill_row(document, node: Tag):
    paragraph = document.add_paragraph()
    label = node.select_one('.skill-label')
    if label:
        run = paragraph.add_run(label.get_text(' ', strip=True))
        run.bold = True
        run.font.size = Pt(BASE_FONT_SIZE)
        run.font.color.rgb = PDF_TEXT
        text = node.get_text(' ', strip=True)
        label_text = label.get_text(' ', strip=True)
        remainder = text.replace(label_text, '', 1).strip()
        if remainder:
            paragraph.add_run(' ' + remainder)
    else:
        _append_inline(node, paragraph, size=BASE_FONT_SIZE)
    _format_paragraph(paragraph, size=BASE_FONT_SIZE, space_after=SKILL_SPACE_AFTER, color=PDF_TEXT)


def _render_education_item(document, node: Tag):
    table = node.select_one('table.edu-header-table')
    if table:
        cells = table.find_all('td')
        if cells:
            degree = cells[0].get_text(' ', strip=True) if len(cells) > 0 else ''
            school = cells[1].get_text(' ', strip=True) if len(cells) > 1 else ''
            date = cells[2].get_text(' ', strip=True) if len(cells) > 2 else ''
            if degree or school or date:
                _add_tabbed_line(
                    document,
                    f'{degree}  {school}'.strip(),
                    date,
                    left_size=BASE_FONT_SIZE,
                    right_size=BASE_FONT_SIZE - 0.5,
                    left_bold=True,
                    left_color=PDF_TEXT,
                    right_color=PDF_DATE,
                    space_after=EDU_TAB_SPACE_AFTER,
                    keep_with_next=True,
                )

    detail = node.select_one('.edu-detail')
    if detail:
        _render_inline_container(document, detail, size=BASE_FONT_SIZE - 0.5, space_after=EDU_DETAIL_SPACE_AFTER)


def _render_lc_row(document, row: Tag, *, space_after: float = 5.0) -> None:
    """
    Render an <li> with one inner <div class="lc-row"> containing a
    <span> (title) and <span class="lc-date"> (year) — matches PDF
    flex layout. Title cell takes 88% of width, date cell 12%, right aligned.
    """
    title_span = row.select_one('.lc-row > span:not(.lc-date)')
    date_span = row.select_one('.lc-row > span.lc-date')
    if not title_span or not date_span:
        return
    title_runs = _collect_inline_runs(title_span, size=BASE_FONT_SIZE - 0.8)
    date_runs = _collect_inline_runs(date_span, size=BASE_FONT_SIZE - 1.0,
                                     color=PDF_DATE)
    if not title_runs and not date_runs:
        return
    _add_borderless_table_row(
        document,
        title_runs or [('', {})],
        date_runs or [('', {})],
        left_size=BASE_FONT_SIZE - 0.8,
        right_size=BASE_FONT_SIZE - 1.0,
        right_width_pct=12.0,
        space_after=space_after,
    )


def _collect_inline_runs(node, *, size: float = BASE_FONT_SIZE,
                          color: Optional[RGBColor] = None):
    """
    Flatten a node's inline content (text, <strong>, <a>) into a list of
    (text, kwargs) tuples for use with _add_borderless_table_row / _add_hyperlink.
    Strips HTML entity decoding — caller must pass plain text strings.
    """
    from bs4 import NavigableString as _NS, Tag as _Tag
    out = []
    if isinstance(node, _NS):
        text = str(node)
        if text.strip():
            out.append((text, {'size': size, 'color': color}))
        return out
    for child in node.children:
        if isinstance(child, _NS):
            text = str(child)
            if text:
                out.append((text, {'size': size, 'color': color}))
        elif isinstance(child, _Tag):
            tag = child.name.lower()
            if tag in {'strong', 'b'}:
                inner = _collect_inline_runs(child, size=size, color=color)
                out.extend([(t, {**k, 'bold': True}) for t, k in inner])
            elif tag == 'a':
                href = (child.get('href') or '').strip()
                text = child.get_text(' ', strip=False)
                if text:
                    if href:
                        out.append((text, {
                            'size': size, 'color': PDF_LINK_BLUE,
                            'hyperlink': href, 'bold': False,
                        }))
                    else:
                        out.append((text, {
                            'size': size, 'color': PDF_LINK_BLUE,
                            'underline': True,
                        }))
            else:
                out.extend(_collect_inline_runs(child, size=size, color=color))
    return out


def _render_job_header(document, node: Tag):
    # node can be the table itself (dispatcher on line ~478) or a wrapper div
    if node.name == 'table':
        table = node
    else:
        table = node.select_one('table.job-header-table')
    if not table:
        return
    cells = table.find_all('td')
    if not cells:
        return

    left = cells[0].get_text(' ', strip=True) if len(cells) > 0 else ''
    right = cells[-1].get_text(' ', strip=True) if len(cells) > 1 else ''
    if left or right:
        _add_tabbed_line(
            document,
            left,
            right,
            left_size=BASE_FONT_SIZE,
            right_size=BASE_FONT_SIZE - 0.5,
            left_bold=True,
            left_color=PDF_TEXT,
            right_color=PDF_DATE,
            space_after=JOB_HEADER_SPACE_AFTER,
            keep_with_next=True,
        )


def _render_node(document, node):
    if isinstance(node, Comment):
        return
    if isinstance(node, NavigableString):
        text = str(node).strip()
        if text:
            _add_text_paragraph(document, text)
        return

    if not isinstance(node, Tag):
        return

    classes = set(node.get('class', []))

    if 'section-title' in classes:
        _render_section_title(document, node.get_text(' ', strip=True))
        return

    if 'cv-header' in classes:
        _render_header(document, node)
        return

    if node.name == 'table' and 'job-header-table' in classes:
        _render_job_header(document, node)
        return

    if 'cv-summary' in classes or 'cv-references' in classes or 'company-alignment' in classes:
        _render_inline_container(
            document,
            node,
            size=BASE_FONT_SIZE,
            italic=('company-alignment' in classes),
            color=PDF_MUTED,
            space_after=SUMMARY_SPACE_AFTER,
        )
        return

    if 'cv-skills' in classes:
        for child in node.find_all(recursive=False):
            _render_node(document, child)
        return

    if 'skill-row' in classes:
        _render_skill_row(document, node)
        return

    if 'edu-item' in classes:
        _render_education_item(document, node)
        return

    if node.name in {'ul', 'ol'}:
        # lc-list is a no-bullet flex list (Licenses / Interests) — render rows
        # as borderless tables so title/date align correctly and there's no bullet glyph.
        if 'lc-list' in classes:
            for li in node.find_all('li', recursive=False):
                if li.select_one('.lc-row'):
                    _render_lc_row(document, li, space_after=LIST_SPACE_AFTER * 3)
                else:
                    paragraph = document.add_paragraph()
                    _append_inline(li, paragraph, size=BASE_FONT_SIZE - 0.8,
                                   color=PDF_TEXT)
                    _format_paragraph(paragraph,
                                       size=BASE_FONT_SIZE - 0.8,
                                       space_after=LIST_SPACE_AFTER * 3,
                                       color=PDF_TEXT)
            return
        _render_list(document, node)
        return

    if any(c in classes for c in {'job', 'career-stage', 'job-employer'}):
        for child in node.find_all(recursive=False):
            _render_node(document, child)
        return

    if 'job-role' in classes or 'stage-header' in classes:
        paragraph = document.add_paragraph()
        strong = node.find('strong')
        if strong:
            label = strong.get_text(' ', strip=True)
            import re
            # 如果 strong 本身包含时间范围，拆分并右对齐时间
            date_match_label = re.search(r'(\b\d{4}(?:\s*[–-]\s*(?:\d{4}|Present|present))\b)', label)
            if date_match_label:
                date_text = date_match_label.group(1)
                left_text = label[:date_match_label.start()].strip(' -—|')
                left_full = left_text
                _add_tabbed_line(
                    document,
                    left_full,
                    date_text,
                    left_size=BASE_FONT_SIZE,
                    right_size=BASE_FONT_SIZE - 0.5,
                    left_bold=True,
                    left_color=PDF_TEXT,
                    right_color=PDF_DATE,
                    space_after=JOB_ROLE_SPACE_AFTER,
                    keep_with_next=True,
                )
                try:
                    p = paragraph._element
                    p.getparent().remove(p)
                except Exception:
                    pass
            else:
                run = paragraph.add_run(label)
                run.bold = True
                run.font.size = Pt(BASE_FONT_SIZE)
                try:
                    run.font.name = FONT_NAME
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                except Exception:
                    pass
                remainder = node.get_text(' ', strip=True).replace(label, '', 1).strip(' -—|')
                if remainder:
                    paragraph.add_run(' ' + remainder)
                paragraph.paragraph_format.keep_with_next = True
        else:
            # 没有 strong 标签的情况：如果结尾包含时间范围，右对齐时间
            text = node.get_text(' ', strip=True)
            import re
            date_match = re.search(r'(\b\d{4}(?:\s*[–-]\s*(?:\d{4}|Present|present))\b)', text)
            if date_match:
                date_text = date_match.group(1)
                left_text = text[:date_match.start()].strip(' -—|')
                _add_tabbed_line(
                    document,
                    left_text,
                    date_text,
                    left_size=BASE_FONT_SIZE,
                    right_size=BASE_FONT_SIZE - 0.5,
                    left_bold=False,
                    left_color=PDF_TEXT,
                    right_color=PDF_DATE,
                    space_after=JOB_ROLE_SPACE_AFTER,
                    keep_with_next=True,
                )
                try:
                    p = paragraph._element
                    p.getparent().remove(p)
                except Exception:
                    pass
            else:
                _append_inline(node, paragraph)
                paragraph.paragraph_format.keep_with_next = True
        _format_paragraph(paragraph, size=BASE_FONT_SIZE, space_after=JOB_ROLE_SPACE_AFTER, color=PDF_TEXT)
        return

    if 'job-tech' in classes or 'stage-tech' in classes or 'stage-focus' in classes or 'employer-desc' in classes:
        _render_inline_container(
            document,
            node,
            size=BASE_FONT_SIZE,
            italic=('stage-focus' in classes),
            color=PDF_MUTED,
            space_after=TECH_SPACE_AFTER,
        )
        return

    if node.name in {'p', 'div'}:
        if node.get_text(' ', strip=True):
            _render_inline_container(document, node, space_after=BASE_SPACE_AFTER)
        return

    if node.name in {'section', 'article'}:
        for child in node.children:
            _render_node(document, child)


def html_to_docx(html_content: str, output_path: str) -> None:
    soup = BeautifulSoup(html_content, 'html.parser')
    document = Document()
    _set_document_defaults(document)

    body = soup.body or soup
    for child in body.children:
        _render_node(document, child)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
