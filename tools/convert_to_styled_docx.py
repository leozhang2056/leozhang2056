from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
import re

src_path = 'templates/CV_Test.docx'
out_path = 'templates/CV_Test_styled_v2.docx'

def is_date_token(s):
    return re.search(r'\b\d{4}\b', s) is not None

src = Document(src_path)
new = Document()

# Layout: set margins similar to PDF
section = new.sections[0]
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)

# Default font and size
normal = new.styles['Normal']
try:
    normal.font.name = 'Inter'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Inter')
except Exception:
    normal.font.name = 'Arial'
normal.font.size = Pt(10)

first_title_done = False
for p in src.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    style_name = p.style.name if p.style else ''
    # Title (first large title)
    if ('Title' in style_name or (len(text.split())<=3 and text==text.upper())) and not first_title_done:
        para = new.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(24)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        first_title_done = True
        continue
    # Contact line
    if '@' in text or '+' in text or 'Portfolio' in text or 'LinkedIn' in text or 'GitHub' in text:
        para = new.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(9)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.paragraph_format.space_after = Pt(6)
        continue
    # Section headings (force uppercase like PDF)
    heading_tokens = ('SUMMARY','KEY SKILLS','EXPERIENCE','EDUCATION','LICENSES','CERTIFICATIONS','PUBLICATIONS')
    if text.strip().upper() in heading_tokens or p.style.name.startswith('Heading'):
        para = new.add_paragraph()
        run = para.add_run(text.strip().upper())
        run.bold = True
        run.font.size = Pt(11)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(2)
        continue
    # Skills or category lines with colon: bold the category
    if ':' in text and len(text) < 200:
        parts = text.split(':',1)
        cat, rest = parts[0].strip(), parts[1].strip()
        para = new.add_paragraph()
        r1 = para.add_run(cat + ': ')
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = para.add_run(rest)
        r2.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(1)
        continue
    # Experience lines with date at end: detect and place date right aligned
    m = re.search(r'(.+?)\s{2,}(\d{4}.*)$', text)
    if not m:
        parts = re.split(r'\s+[-–—]\s+', text)
        if len(parts) >= 2 and is_date_token(parts[-1]):
            left = ' - '.join(parts[:-1])
            right = parts[-1]
            m = (left, right)
    if m:
        if isinstance(m, tuple):
            left, right = m
        else:
            left, right = m.group(1).strip(), m.group(2).strip()
        para = new.add_paragraph()
        tab_pos = Inches(6.6)
        para.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        run = para.add_run(left + '\t' + right)
        run.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(1)
        continue
    # Default body paragraph
    para = new.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(1)

new.save(out_path)
print('WROTE', out_path)
