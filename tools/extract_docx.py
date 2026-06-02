from docx import Document
from datetime import datetime
import sys

path = 'templates/CV_Test.docx'
try:
    doc = Document(path)
except Exception as e:
    print('ERROR_OPEN', e)
    sys.exit(1)

props = doc.core_properties
metadata = {
    'title': props.title,
    'author': props.author,
    'last_modified_by': props.last_modified_by,
    'created': str(props.created) if props.created else None,
    'modified': str(props.modified) if props.modified else None,
    'comments': props.comments,
}

lines = []
lines.append('---METADATA---')
for k, v in metadata.items():
    lines.append(f'{k}: {v}')

lines.append('\n---PARAGRAPHS---')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    style = p.style.name if p.style else ''
    lines.append(f'[{i}] ({style}) {text}')

lines.append('\n---TABLES---')
for ti, table in enumerate(doc.tables):
    lines.append(f'-- Table {ti} --')
    for r in table.rows:
        row_cells = [c.text.replace('\n', ' ') for c in r.cells]
        lines.append(' | '.join(row_cells))

out_path = 'outputs/CV_Test_extracted.txt'
with open(out_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print('WROTE', out_path)
print('Paragraphs:', len(doc.paragraphs))
print('Tables:', len(doc.tables))
print('Metadata:', metadata)